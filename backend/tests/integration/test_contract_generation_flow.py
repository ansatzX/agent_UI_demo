# backend/tests/integration/test_contract_generation_flow.py
import pytest
from pathlib import Path
from backend.src.services.doc_generator import DocGenerator
from backend.src.services.tool_registry import ToolRegistry
from backend.src.services.react_agent import ReActAgent
from backend.src.services.llm_service import LLMService
from backend.src.services.tools.generate_document import GenerateDocumentTool
from backend.src.services.tools.show_form import ShowFormTool

@pytest.mark.asyncio
async def test_full_contract_generation_flow(tmp_path):
    """测试完整的合同生成流程"""
    # 1. 初始化服务
    doc_generator = DocGenerator(tmp_path, tmp_path / "output")

    tool_registry = ToolRegistry()
    tool_registry.register(GenerateDocumentTool(doc_generator))
    tool_registry.register(ShowFormTool())

    llm_service = LLMService()
    agent = ReActAgent(llm_service, tool_registry)

    # 2. 创建测试模板
    from docx import Document
    template_file = tmp_path / "template_1.docx"
    doc = Document()
    doc.add_paragraph("甲方：{{party_a}}")
    doc.add_paragraph("乙方：{{party_b}}")
    doc.save(str(template_file))

    # 3. 运行 Agent
    result = await agent.run(
        "帮我生成一份合同，甲方是香港中文大学，乙方是XX公司",
        {}
    )

    # 4. 验证结果
    assert result.message is not None
