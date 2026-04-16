# backend/tests/integration/test_file_upload_fix.py
"""文件上传修复端到端测试"""
import pytest
from pathlib import Path
from backend.src.services.mcp_client import MCPClient
from backend.src.services.file_service import FileService
from backend.src.services.react_agent import ReActAgent
from backend.src.services.llm_service import LLMService
from backend.src.services.tool_registry import ToolRegistry


@pytest.mark.asyncio
async def test_file_upload_and_mcp_processing():
    """测试文件上传和 MCP 处理完整流程"""
    # 1. 初始化
    mcp_client = MCPClient()
    await mcp_client.connect()

    file_service = FileService(mcp_client)

    # 2. 创建测试文档
    from docx import Document
    test_file = Path("/tmp/test_contract.docx")
    doc = Document()
    doc.add_paragraph("测试合同内容")
    doc.add_paragraph("甲方：测试公司")
    doc.add_paragraph("乙方：测试客户")
    doc.add_paragraph("合同金额：100万元")
    doc.save(str(test_file))

    # 3. 处理文件
    result = await file_service.process_uploaded_file(
        test_file,
        "test_contract.docx"
    )

    # 4. 验证结果
    assert result["success"] is True
    assert "paragraphs" in result
    assert len(result["paragraphs"]) > 0
    assert "full_text" in result
    assert "测试合同内容" in result["full_text"]

    # 5. 清理
    await mcp_client.close()
    test_file.unlink()


@pytest.mark.asyncio
async def test_agent_with_uploaded_file():
    """测试 Agent 能否识别上传的文件"""
    # 模拟上传文件
    uploaded_file = {
        "filename": "test.docx",
        "original_filename": "测试合同.docx",
        "size": 1024,
        "content": {
            "success": True,
            "paragraphs": ["测试段落1", "测试段落2"],
            "tables": [],
            "full_text": "这是测试内容，包含重要信息。"
        }
    }

    # 测试 Agent
    llm = LLMService()
    tools = ToolRegistry()
    agent = ReActAgent(llm, tools)

    # 构建对话
    messages = agent._build_conversation(
        "看看这个文件",
        [],
        {"uploaded_file": uploaded_file}
    )

    # 验证消息包含文件内容
    assert "测试合同.docx" in messages[0]["content"]
    assert "测试内容" in messages[0]["content"]
    assert "段落数: 2" in messages[0]["content"]


@pytest.mark.asyncio
async def test_agent_with_failed_file():
    """测试 Agent 处理解析失败的文件"""
    # 模拟解析失败的文件
    uploaded_file = {
        "filename": "failed.docx",
        "original_filename": "失败文件.docx",
        "size": 1024,
        "content": {
            "success": False,
            "error": "Invalid request parameters"
        }
    }

    # 测试 Agent
    llm = LLMService()
    tools = ToolRegistry()
    agent = ReActAgent(llm, tools)

    # 构建对话
    messages = agent._build_conversation(
        "看看这个文件",
        [],
        {"uploaded_file": uploaded_file}
    )

    # 验证消息包含错误信息
    assert "失败文件.docx" in messages[0]["content"]
    assert "解析失败" in messages[0]["content"]
