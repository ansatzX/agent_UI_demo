# backend/src/services/tools/save_document.py
from datetime import datetime
from pathlib import Path
from urllib.parse import quote
import uuid

from .base import Tool
from .base import ToolResult


class SaveDocumentTool(Tool):
    """保存内容为 Word 文档"""

    name = "save_document"
    description = "将文本内容保存为 Word 文档（.docx）并提供下载链接"
    parameters = {
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "文档标题"},
            "content": {
                "type": "string",
                "description": "文档内容（支持 Markdown 格式）",
            },
            "document_type": {
                "type": "string",
                "description": "文档类型（如：项目报告、新闻稿、文章等）",
                "default": "文章",
            },
        },
        "required": ["title", "content"],
    }

    def __init__(self, uploads_dir: Path = None):
        self.uploads_dir = uploads_dir or Path("uploads")
        self.uploads_dir.mkdir(exist_ok=True)

    async def execute(
        self, title: str, content: str, document_type: str = "文章"
    ) -> ToolResult:
        """生成 Word 文档"""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
            from docx.shared import Pt

            # 创建文档
            doc = Document()

            # 添加标题
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加文档类型
            type_para = doc.add_paragraph(f"文档类型：{document_type}")
            type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加日期
            date_para = doc.add_paragraph(
                f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}"
            )
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加分隔线
            doc.add_paragraph("_" * 60)

            # 处理内容：按段落分割
            paragraphs = content.strip().split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # 检测标题（以 # 开头的 Markdown 标题）
                if para_text.startswith("# "):
                    heading_text = para_text[2:].strip()
                    doc.add_heading(heading_text, level=1)
                elif para_text.startswith("## "):
                    heading_text = para_text[3:].strip()
                    doc.add_heading(heading_text, level=2)
                elif para_text.startswith("### "):
                    heading_text = para_text[4:].strip()
                    doc.add_heading(heading_text, level=3)
                elif para_text.startswith("#### "):
                    heading_text = para_text[5:].strip()
                    doc.add_heading(heading_text, level=4)
                else:
                    # 普通段落
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.first_line_indent = Inches(0.3)
                    para.paragraph_format.space_after = Pt(12)

            # 生成唯一文件名
            prefix = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
            safe_title = "".join(
                c for c in title if c.isalnum() or c in (" ", "-", "_")
            )[:30]
            filename = f"{prefix}_{safe_title}.docx"
            output_path = self.uploads_dir / filename

            # 保存文档
            doc.save(str(output_path))

            return ToolResult(
                success=True,
                output={
                    "file_path": str(output_path),
                    "filename": filename,
                    "display_name": f"{title}.docx",
                    "download_url": f"/api/files/download/{quote(filename)}",
                },
            )
        except Exception as e:
            return ToolResult(
                success=False, output={}, error=f"文档生成失败: {str(e)}"
            )
