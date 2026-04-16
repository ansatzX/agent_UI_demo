from pathlib import Path
from typing import Optional, Dict, Any
from ..config import settings
import uuid
import re
from datetime import datetime
import json
import logging
import asyncio

logger = logging.getLogger(__name__)

class FileService:
    def __init__(self):
        self.uploads_dir = settings.project_root / "uploads"
        self.uploads_dir.mkdir(exist_ok=True)
        self.processed_dir = self.uploads_dir / "processed"
        self.processed_dir.mkdir(exist_ok=True)

    def generate_unique_filename(self, original_filename: str) -> str:
        """生成唯一的文件名，避免特殊字符问题"""
        # 提取文件扩展名
        ext = Path(original_filename).suffix.lower()

        # 清理原始文件名，移除特殊字符
        clean_name = re.sub(r'[^\w\u4e00-\u9fff]', '_', Path(original_filename).stem)
        clean_name = clean_name.strip('_')[:50]  # 限制长度

        # 生成唯一ID
        unique_id = datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:8]

        # 组合成唯一文件名
        unique_filename = f"{unique_id}_{clean_name}{ext}"

        return unique_filename

    def save_upload_file(self, file_content: bytes, original_filename: str) -> tuple[Path, str]:
        """保存上传的文件，返回文件路径和唯一文件名"""
        unique_filename = self.generate_unique_filename(original_filename)
        file_path = self.uploads_dir / unique_filename

        with open(file_path, "wb") as f:
            f.write(file_content)

        return file_path, unique_filename

    async def process_uploaded_file(
        self,
        file_path: Path,
        unique_filename: str
    ) -> Dict[str, Any]:
        """处理上传的文件（使用简单解析，不依赖 MCP）"""
        logger.info(f"处理文件: {file_path} ({unique_filename})")

        try:
            # 检查文件格式
            ext = file_path.suffix.lower()

            if ext == '.docx':
                # 使用 asyncio.to_thread 避免阻塞事件循环
                result = await asyncio.to_thread(self._parse_docx_simple, file_path)
            elif ext == '.doc':
                logger.warning(f".doc 格式暂不支持: {file_path}")
                return {
                    "success": False,
                    "error": "暂不支持 .doc 格式，请上传 .docx 格式文件"
                }
            else:
                return {
                    "success": False,
                    "error": f"不支持的文件格式: {ext}"
                }

            logger.info(f"文件解析完成: 段落={len(result.get('paragraphs', []))}, "
                       f"表格={len(result.get('tables', []))}, "
                       f"文本长度={len(result.get('full_text', ''))}")

            return result

        except Exception as e:
            logger.error(f"文件处理异常: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e)
            }

    def _parse_docx_simple(self, file_path: Path) -> Dict[str, Any]:
        """使用 python-docx 简单解析 .docx 文件"""
        try:
            from docx import Document

            doc = Document(str(file_path))

            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # 组合完整文本
            full_text = '\n\n'.join(paragraphs)

            return {
                "success": True,
                "paragraphs": paragraphs,
                "tables": tables,
                "full_text": full_text,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "character_count": len(full_text)
            }

        except Exception as e:
            logger.error(f"解析 docx 文件失败: {e}", exc_info=True)
            return {
                "success": False,
                "error": f"文件解析失败: {str(e)}"
            }

    def get_file_info(self, filename: str) -> Optional[Dict[str, Any]]:
        """获取文件信息"""
        file_path = self.uploads_dir / filename
        if not file_path.exists():
            return None

        return {
            "filename": filename,
            "path": str(file_path),
            "size": file_path.stat().st_size
        }

    def list_uploaded_files(self) -> list:
        """列出所有上传的文件"""
        files = []
        for file_path in self.uploads_dir.glob("*.docx"):
            files.append({
                "filename": file_path.name,
                "size": file_path.stat().st_size
            })
        for file_path in self.uploads_dir.glob("*.doc"):
            files.append({
                "filename": file_path.name,
                "size": file_path.stat().st_size
            })
        return files
