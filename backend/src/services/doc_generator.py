# backend/src/services/doc_generator.py
import logging
from pathlib import Path
from typing import Dict

from docx import Document
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)


class DocGenerator:
    """文档生成器"""

    def __init__(self, template_dir: Path, output_dir: Path):
        self.template_dir = template_dir
        self.output_dir = output_dir
        self.output_dir.mkdir(exist_ok=True)

    async def fill_template(
        self, template_id: int, fields: Dict[str, str], output_filename: str
    ) -> Path:
        """
        填充模板字段生成新文档（带错误处理）

        Args:
            template_id: 模板 ID（从数据库加载模板）
            fields: 字段名 -> 值的映射
            output_filename: 输出文件名

        Returns:
            生成的文档路径
        """
        try:
            # 加载模板（这里简化处理，实际应从数据库加载）
            template_path = self.template_dir / f"template_{template_id}.docx"

            if not template_path.exists():
                logger.error(f"模板文件不存在: {template_path}")
                raise FileNotFoundError(f"模板文件不存在: {template_path}")

            if not fields:
                logger.error("字段不能为空")
                raise ValueError("字段不能为空")

            # 使用 docxtpl 填充模板
            doc = DocxTemplate(str(template_path))

            # 渲染模板
            doc.render(fields)

            # 保存文档
            output_path = self.output_dir / output_filename
            doc.save(str(output_path))

            logger.info(f"文档生成成功: {output_path}")
            return output_path

        except FileNotFoundError as e:
            logger.error(f"模板文件错误: {e}")
            raise
        except PermissionError as e:
            logger.error(f"文件权限错误: {e}")
            raise
        except Exception as e:
            logger.error(f"文档生成失败: {e}", exc_info=True)
            raise

    async def fill_template_simple(
        self, template_path: Path, fields: Dict[str, str], output_path: Path
    ) -> Path:
        """
        简单的模板填充（使用 python-docx，带错误处理）

        适用于没有使用 docxtpl 标记的文档
        """
        try:
            if not template_path.exists():
                logger.error(f"模板文件不存在: {template_path}")
                raise FileNotFoundError(f"模板文件不存在: {template_path}")

            if not fields:
                logger.error("字段不能为空")
                raise ValueError("字段不能为空")

            doc = Document(str(template_path))

            # 替换段落中的占位符
            for paragraph in doc.paragraphs:
                for key, value in fields.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            placeholder, value
                        )

            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in fields.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(
                                    placeholder, value
                                )

            # 保存文档
            doc.save(str(output_path))

            logger.info(f"文档生成成功: {output_path}")
            return output_path

        except FileNotFoundError as e:
            logger.error(f"模板文件错误: {e}")
            raise
        except PermissionError as e:
            logger.error(f"文件权限错误: {e}")
            raise
        except Exception as e:
            logger.error(f"文档生成失败: {e}", exc_info=True)
            raise
