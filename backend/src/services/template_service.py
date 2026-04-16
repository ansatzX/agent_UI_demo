from sqlmodel import Session, select
from typing import List, Optional
from docx import Document
import io
from ..models.template import Template
from ..schemas.template import FieldInfo
from .llm_service import LLMService


class TemplateService:
    def __init__(self, session: Session, llm_service: LLMService):
        self.session = session
        self.llm_service = llm_service

    async def create_template(
        self,
        name: str,
        type: str,
        file_content: bytes,
        description: Optional[str] = None
    ) -> Template:
        template = Template(
            name=name,
            type=type,
            description=description,
            content=file_content
        )

        fields = await self.analyze_template_fields(file_content)
        if fields:
            import json
            template.field_config = json.dumps([f.model_dump() for f in fields], ensure_ascii=False)

        self.session.add(template)
        self.session.commit()
        self.session.refresh(template)
        return template

    async def analyze_template_fields(self, file_content: bytes) -> List[FieldInfo]:
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n".join(text_parts)

            field_dicts = await self.llm_service.analyze_template(full_text[:8000])

            fields = []
            for fd in field_dicts:
                fields.append(FieldInfo(
                    name=fd.get("name", ""),
                    label=fd.get("label", ""),
                    placeholder=fd.get("placeholder"),
                    field_type=fd.get("field_type", "text"),
                    group=fd.get("group"),
                    required=fd.get("required", True)
                ))

            return fields
        except Exception as e:
            print(f"Error analyzing template: {e}")
            return []

    def get_template(self, template_id: int) -> Optional[Template]:
        return self.session.get(Template, template_id)

    def list_templates(self, type_filter: Optional[str] = None) -> List[Template]:
        query = select(Template).order_by(Template.created_at.desc())
        if type_filter:
            query = query.where(Template.type == type_filter)
        return self.session.exec(query).all()
